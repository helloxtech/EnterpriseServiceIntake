#!/usr/bin/env python3
"""Build the Enterprise Service Intake user manual DOCX/PDF artifacts."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from build_architecture_design_doc import (
    BLUE,
    DARK_BLUE,
    DARK_GREY,
    GREY,
    INK,
    MID_GREY,
    ROOT,
    SUBMISSION_DIR,
    add_bullets,
    add_key_value_table,
    add_numbered,
    add_section_title,
    add_standard_table,
    add_subtitle,
    configure_styles,
    create_pdf_styles,
    pdf_bullets,
    pdf_color,
    pdf_heading,
    pdf_key_value_table,
    pdf_numbers,
    pdf_table,
    pp,
    rgb,
)


DOCX_PATH = SUBMISSION_DIR / "Enterprise_ServiceIntake_User_Manual_ForrestZhang_v2.docx"
PDF_PATH = SUBMISSION_DIR / "Enterprise_ServiceIntake_User_Manual_ForrestZhang_v2.pdf"
SOURCE_PATH = ROOT / "docs" / "manual" / "user-manual.md"


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Enterprise Service Intake")
    run.font.size = Pt(25)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run("User Manual - V2")
    sub_run.font.size = Pt(18)
    sub_run.bold = True
    sub_run.font.color.rgb = rgb(BLUE)

    add_key_value_table(
        doc,
        [
            ("Candidate", "Forrest Zhang"),
            ("Prepared For", "Mitacs hiring and technical review team"),
            ("Audience", "Reviewers, portal customers, coordinators, managers, and administrators"),
            ("Environment", "https://mitacs.crm.dynamics.com/"),
            ("Portal", "https://enterprise-service-intake-hellox.powerappsportals.com"),
            ("Status", "Candidate user manual - V2"),
        ],
        widths=(1.8, 4.9),
    )

    add_subtitle(doc, "How To Use This Manual")
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Body Text"]
    paragraph.add_run(
        "This guide explains the live review path: customer portal submission, required-file handling, internal coordinator review, routing matrix administration, approval/ERP sync, and troubleshooting evidence."
    )


def add_quick_access(doc: Document) -> None:
    add_section_title(doc, "Quick Access")
    add_standard_table(
        doc,
        ["Area", "URL / Location", "Purpose"],
        [
            ["Power Pages site", "https://enterprise-service-intake-hellox.powerappsportals.com", "External customer intake, drafts, required-file upload, and final submission"],
            ["Model-driven app", "https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f", "Internal coordinator, manager, configuration, and monitoring experience"],
            ["Maker solution", "https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions", "Solution components, flows, tables, forms, views, and web resources"],
            ["Hidden ERP console", "https://hellox.ca/esi/", "View mock ERP sync attempts, returned external IDs, and failure-path evidence"],
        ],
        [1.25, 2.85, 2.65],
        font_size=7,
    )
    add_bullets(
        doc,
        [
            "Reviewer passwords are not stored in Git.",
            "The administrator should share passwords out of band and rotate them after the interview.",
        ],
    )


def add_portal(doc: Document) -> None:
    add_section_title(doc, "External Customer Portal")
    add_subtitle(doc, "Create Or Resume A Request")
    add_numbered(
        doc,
        [
            "Open the Power Pages site and sign in.",
            "Select New service request to start a new intake.",
            "Use My requests to resume an existing draft or review submitted requests.",
            "Use View details on a request card to inspect the confirmation number, status, routing estimate, request text, and supporting-file requirement.",
            "Use Save for later if you want to leave the intake before final submission.",
        ],
    )
    add_bullets(
        doc,
        [
            "Draft requests remain in Draft and do not trigger the applicant confirmation email.",
            "Sign-in is required before creating, saving, resuming, or submitting a request so each request and file upload stays linked to the correct portal account.",
            "Portal users do not see internal-only notes, error logs, approval details, or integration payloads.",
        ],
    )

    add_subtitle(doc, "Request Details And Impact")
    add_numbered(
        doc,
        [
            "Enter request title, service category, and description.",
            "Select Impact level and Urgency.",
            "Enter the business impact.",
            "Watch the response estimate update without a full page reload.",
        ],
    )
    add_bullets(
        doc,
        [
            "The estimate shows team, response target, Mitacs review, and required-file status.",
            "Missing required values are shown inline before the user can continue.",
        ],
    )

    add_subtitle(doc, "Supporting Files")
    add_standard_table(
        doc,
        ["Scenario", "User Action", "Expected Result"],
        [
            ["Files required", ["Continue to Files", "Upload at least one file", "Wait for uploaded-file count", "Continue to Review"], "Review remains disabled until at least one file exists."],
            ["Files optional", ["Continue through Files", "Submit from Review", "Add optional files after submission if needed"], "The request can be submitted without an upload."],
        ],
        [1.3, 3.0, 2.45],
        font_size=8,
    )

    add_subtitle(doc, "Submit And Confirm")
    add_numbered(
        doc,
        [
            "Review the request summary.",
            "Select Submit service request.",
            "Capture the confirmation number in the format SR-yyyyMMdd-######.",
            "Return to My requests to view the submitted request.",
            "Use View details if reviewers want to inspect the submitted request from the portal user's perspective.",
        ],
    )


def add_internal_app(doc: Document) -> None:
    add_section_title(doc, "Internal Coordinator App")
    add_standard_table(
        doc,
        ["Navigation Group", "Use"],
        [
            ["Intake Work", "Service Requests and Service Request Evidence Reviews"],
            ["Routing Configuration", "Routing Matrix, Departments, SLA Policies, and Service Categories"],
            ["Monitoring", "System Error Logs and External Sync Logs"],
        ],
        [2.1, 4.65],
        font_size=9,
    )
    add_subtitle(doc, "Role-Based Dashboards")
    add_standard_table(
        doc,
        ["User", "Required Roles", "Expected Dashboards"],
        [
            ["agent@hellosmart.ca", "Basic User; ESI Service Coordinator", "Operations Dashboard and Monitoring Dashboard"],
            ["manager@hellosmart.ca", "Basic User; ESI Approval Manager; Approvals User", "Approval Dashboard and Monitoring Dashboard"],
            ["forrest@hellosmart.ca", "System Administrator or System Customizer", "All dashboards for administration and review"],
        ],
        [1.65, 2.35, 2.75],
        font_size=8,
    )
    add_bullets(
        doc,
        [
            "Approvals User supports Power Automate approval records; it does not control the ESI dashboard visibility.",
            "If a recently changed dashboard list looks stale, sign out and back in or open a fresh browser session.",
        ],
    )
    add_subtitle(doc, "Review A Service Request")
    add_numbered(
        doc,
        [
            "Open Intake Work > Service Requests.",
            "Use Active Service Requests to scan confirmation number, customer, category, impact, urgency, lifecycle, department, SLA due date, approval, ERP sync, and created date.",
            "Open a request and review customer details, triage inputs, routing/SLA, approval/ERP sync, and resolution fields.",
            "Use the PCF status indicator as a compact visual summary of severity, SLA, approval, and sync state.",
        ],
    )
    add_subtitle(doc, "Review Documents")
    add_numbered(
        doc,
        [
            "Open a Service Request.",
            "Open the Documents tab.",
            "Use the SharePoint Documents grid to view files uploaded through Power Pages document management.",
            "Use the SR Evidence Reviews subgrid to record internal review status, file URL, document type, verification status, and notes.",
        ],
    )
    add_subtitle(doc, "Close A Critical Request")
    add_numbered(
        doc,
        [
            "Try to close a critical request without internal resolution notes and accepted evidence.",
            "Confirm that the plugin blocks the close.",
            "Add internal resolution notes.",
            "Add or update an accepted Service Request Evidence Review row with a SharePoint file URL.",
            "Close the request again.",
        ],
    )


def add_routing_matrix(doc: Document) -> None:
    add_section_title(doc, "Routing Matrix Administration")
    add_numbered(
        doc,
        [
            "Open Routing Configuration > Routing Matrix.",
            "Select a service category tab.",
            "Review summary cards for department, SLA distribution, manager review, and required documents.",
            "Find the Impact row and Urgency column to adjust.",
            "Change Department or SLA from the dropdown.",
            "Use toggles for Review and Docs.",
            "Wait for the saved status.",
        ],
    )
    add_bullets(
        doc,
        [
            "The page updates the same Dataverse Routing / SLA Rule rows used by the portal preview and routing plugin.",
            "Routing Matrix presents the 80 active exact-match cells for 5 categories x 4 impact levels x 4 urgency levels.",
            "The raw Dataverse table also includes one active Generic fallback - unmatched request, so the clean final rule set is 81 active rules total.",
            "Runtime logic filters to active rules and uses the fallback only when no exact active match exists.",
            "Select the rule name inside a matrix cell to open the underlying Dataverse record.",
            "If you change a rule during the live demo, change it back immediately after showing the saved status.",
        ],
    )


def add_monitoring(doc: Document) -> None:
    add_section_title(doc, "Approval, ERP Sync, And Monitoring")
    add_subtitle(doc, "Manager Approval And ERP Sync")
    add_numbered(
        doc,
        [
            "Open an approval-required request.",
            "Confirm approval status is pending.",
            "Open Power Automate flow ESI - Approval and ERP Sync.",
            "Review the Try scope, approval action, OAuth token request, protected HTTP POST, Service Request update, External Sync Log creation, and Catch scope.",
            "Use Approval records and flow run history as evidence if email delivery is restricted.",
        ],
    )
    add_subtitle(doc, "Troubleshooting")
    add_standard_table(
        doc,
        ["Surface", "Use"],
        [
            ["System Error Logs", "Plugin, portal, flow, approval, and integration failures with source, stage, message, correlation ID, related request, and resolved status."],
            ["External Sync Logs", "ERP endpoint calls, HTTP status, external ID, payload snapshot, and timestamps."],
            ["HelloX ERP Dashboard", "Open https://hellox.ca/esi/ to view received mock ERP requests, returned external IDs, accepted/failed counts, and breakdowns by status, priority, and department."],
            ["Flow Run Histories", "Confirmation email, approval, ERP sync, and Catch-path evidence."],
        ],
        [1.7, 5.05],
        font_size=8,
    )


def add_accounts_notes(doc: Document) -> None:
    add_section_title(doc, "Demo Accounts And Notes")
    add_standard_table(
        doc,
        ["Account", "Purpose", "Sign-In Location", "Credential Note"],
        [
            ["agent@hellosmart.ca", "Internal service coordinator", "Model-driven app / Power Platform environment", "Uses the temporary demo password shared separately by the administrator"],
            ["manager@hellosmart.ca", "Approval manager", "Model-driven app / Power Platform environment", "Uses the temporary demo password shared separately by the administrator"],
            ["forrest@hellosmart.ca", "Admin/customizer reviewer", "Maker portal, model-driven app, and Power Pages administration", "Admin password is shared separately by the administrator"],
        ],
        [1.45, 1.65, 1.85, 1.8],
        font_size=7,
    )
    add_bullets(
        doc,
        [
            "Do not commit passwords or OAuth secrets.",
            "Power Platform environment: https://mitacs.crm.dynamics.com/",
            "Maker solution: https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions",
            "Model-driven app: https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f",
            "Power Pages site: https://enterprise-service-intake-hellox.powerappsportals.com",
            "Hidden ERP console: https://hellox.ca/esi/",
            "The portal is private for the interview tenant.",
            "The raw Routing / SLA Rule table still exists for Dataverse administration, but normal rule maintenance should use Routing Matrix.",
            "The clean final routing table should contain 81 active rules: 80 exact-match matrix rows plus one active generic fallback.",
            "No final solution export should be produced until the administrator requests it.",
        ],
    )


def manual_pdf_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFillColorRGB(1, 1, 1)
    canvas.rect(0, 0, letter[0], letter[1], stroke=0, fill=1)
    canvas.setStrokeColor(pdf_color(BLUE))
    canvas.setLineWidth(0.6)
    canvas.line(0.75 * inch, 10.33 * inch, 7.75 * inch, 10.33 * inch)
    canvas.setFillColor(pdf_color(MID_GREY))
    canvas.setFont("Helvetica", 7.5)
    canvas.drawRightString(7.75 * inch, 10.43 * inch, "Enterprise Service Intake | User Manual")
    canvas.drawString(0.75 * inch, 0.45 * inch, "Prepared for Mitacs hiring review - May 2026")
    canvas.drawRightString(7.75 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def manual_pdf_styles() -> dict[str, ParagraphStyle]:
    styles = create_pdf_styles()
    styles["Title"].fontSize = 24
    styles["Subtitle"].fontSize = 15
    styles["Note"] = ParagraphStyle(
        "ESINote",
        parent=styles["Body"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11.5,
        textColor=pdf_color(DARK_GREY),
        leftIndent=0,
        spaceBefore=2,
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    return styles


def build_pdf() -> None:
    styles = manual_pdf_styles()
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.82 * inch,
        bottomMargin=0.72 * inch,
        title="Enterprise Service Intake User Manual",
        author="Forrest Zhang",
        subject="Power Platform Senior Developer Take-Home Case",
    )
    story: list = []
    story.append(pp("Enterprise Service Intake", styles["Title"]))
    story.append(pp("User Manual - V2", styles["Subtitle"]))
    pdf_key_value_table(
        story,
        [
            ("Candidate", "Forrest Zhang"),
            ("Prepared For", "Mitacs hiring and technical review team"),
            ("Audience", "Reviewers, portal customers, coordinators, managers, and administrators"),
            ("Environment", "https://mitacs.crm.dynamics.com/"),
            ("Portal", "https://enterprise-service-intake-hellox.powerappsportals.com"),
            ("Status", "Candidate user manual - V2"),
        ],
        styles,
    )
    story.append(pp("This guide explains the live review path: customer portal submission, required-file handling, internal coordinator review, routing matrix administration, approval/ERP sync, and troubleshooting evidence.", styles["Body"]))

    pdf_heading(story, "Quick Access", styles)
    pdf_table(
        story,
        ["Area", "URL / Location", "Purpose"],
        [
            ["Power Pages site", "https://enterprise-service-intake-hellox.powerappsportals.com", "External customer intake, drafts, required-file upload, and final submission"],
            ["Model-driven app", "https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f", "Internal coordinator, manager, configuration, and monitoring experience"],
            ["Maker solution", "https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions", "Solution components, flows, tables, forms, views, and web resources"],
            ["Hidden ERP console", "https://hellox.ca/esi/", "View mock ERP sync attempts, returned external IDs, and failure-path evidence"],
        ],
        [1.15, 2.9, 2.7],
        styles,
    )
    pdf_bullets(story, ["Reviewer passwords are not stored in Git.", "The administrator should share passwords out of band and rotate them after the interview."], styles)

    pdf_heading(story, "External Customer Portal", styles)
    pdf_heading(story, "Create Or Resume A Request", styles, level=2)
    pdf_numbers(story, ["Open the Power Pages site and sign in.", "Select New service request.", "Use My requests to resume a draft or review submitted requests.", "Use View details on a request card to inspect portal-safe request details.", "Use Save for later if you want to leave before final submission."], styles)
    pdf_bullets(story, ["Draft requests remain in Draft and do not trigger the applicant confirmation email.", "Sign-in is required before creating, saving, resuming, or submitting a request so each request and file upload stays linked to the correct portal account.", "Portal users do not see internal-only notes, error logs, approval details, or integration payloads."], styles)
    pdf_heading(story, "Request Details, Impact, And Files", styles, level=2)
    pdf_numbers(story, ["Enter request title, category, and description.", "Select Impact level and Urgency.", "Watch the response estimate update without a full page reload.", "Upload at least one file when documentation is required.", "Submit from Review and capture the confirmation number."], styles)
    pdf_table(
        story,
        ["Scenario", "Expected Result"],
        [
            ["Files required", "Review remains disabled until at least one file exists."],
            ["Files optional", "The request can be submitted without an upload; optional files can be added after submission."],
        ],
        [1.5, 5.25],
        styles,
    )

    pdf_heading(story, "Internal Coordinator App", styles)
    pdf_table(
        story,
        ["Navigation Group", "Use"],
        [
            ["Intake Work", "Service Requests and Service Request Evidence Reviews"],
            ["Routing Configuration", "Routing Matrix, Departments, SLA Policies, and Service Categories"],
            ["Monitoring", "System Error Logs and External Sync Logs"],
        ],
        [2.0, 4.75],
        styles,
    )
    pdf_heading(story, "Role-Based Dashboards", styles, level=2)
    pdf_table(
        story,
        ["User", "Required Roles", "Expected Dashboards"],
        [
            ["agent@hellosmart.ca", "Basic User; ESI Service Coordinator", "Operations Dashboard and Monitoring Dashboard"],
            ["manager@hellosmart.ca", "Basic User; ESI Approval Manager; Approvals User", "Approval Dashboard and Monitoring Dashboard"],
            ["forrest@hellosmart.ca", "System Administrator or System Customizer", "All dashboards for administration and review"],
        ],
        [1.55, 2.35, 2.85],
        styles,
    )
    pdf_bullets(
        story,
        [
            "Approvals User supports Power Automate approval records; it does not control ESI dashboard visibility.",
            "If a recently changed dashboard list looks stale, sign out and back in or open a fresh browser session.",
        ],
        styles,
    )
    pdf_numbers(story, ["Open Intake Work > Service Requests.", "Use Active Service Requests to scan operational columns.", "Open a request and review customer, triage, routing/SLA, approval/ERP sync, and resolution fields.", "Open Documents to view SharePoint files and SR Evidence Reviews."], styles)
    pdf_bullets(story, ["Critical closure is blocked until internal resolution notes and accepted evidence with a SharePoint file URL exist."], styles)

    story.append(PageBreak())
    pdf_heading(story, "Routing Matrix Administration", styles)
    pdf_numbers(story, ["Open Routing Configuration > Routing Matrix.", "Select a service category tab.", "Review summary cards for department, SLA distribution, manager review, and required documents.", "Edit Department or SLA from dropdowns.", "Use toggles for Review and Docs.", "Wait for the saved status."], styles)
    pdf_bullets(story, ["The page updates the same Dataverse rules used by the portal preview and routing plugin.", "Routing Matrix presents the 80 active exact-match cells.", "The raw Dataverse table also includes one active generic fallback, so the clean final rule set is 81 active rules total.", "Runtime logic filters active rules and uses the fallback only when no exact active match exists.", "Select a rule name to open the underlying Dataverse record.", "If you change a rule during the live demo, change it back immediately."], styles)

    pdf_heading(story, "Approval, ERP Sync, And Monitoring", styles)
    pdf_numbers(story, ["Open an approval-required request.", "Open Power Automate flow ESI - Approval and ERP Sync.", "Review the Try scope, approval action, OAuth token request, protected HTTP POST, Service Request update, External Sync Log, and Catch scope.", "Use Approval records and flow run history as evidence if email delivery is restricted."], styles)
    pdf_table(
        story,
        ["Surface", "Use"],
        [
            ["System Error Logs", "Plugin, portal, flow, approval, and integration failures."],
            ["External Sync Logs", "ERP endpoint calls, HTTP status, external ID, payload snapshot, and timestamps."],
            ["HelloX ERP Dashboard", "Open https://hellox.ca/esi/ to view received mock ERP requests, returned external IDs, accepted/failed counts, and breakdowns by status, priority, and department."],
            ["Flow Run Histories", "Confirmation email, approval, ERP sync, and Catch-path evidence."],
        ],
        [1.65, 5.1],
        styles,
    )
    pdf_heading(story, "Demo Accounts And Notes", styles)
    pdf_table(
        story,
        ["Account", "Purpose", "Sign-In Location", "Credential Note"],
        [
            ["agent@hellosmart.ca", "Internal service coordinator", "Model-driven app / Power Platform environment", "Uses the temporary demo password shared separately by the administrator"],
            ["manager@hellosmart.ca", "Approval manager", "Model-driven app / Power Platform environment", "Uses the temporary demo password shared separately by the administrator"],
            ["forrest@hellosmart.ca", "Admin/customizer reviewer", "Maker portal, model-driven app, and Power Pages administration", "Admin password is shared separately by the administrator"],
        ],
        [1.35, 1.45, 1.75, 2.2],
        styles,
    )
    pdf_bullets(story, ["Do not commit passwords or OAuth secrets.", "Power Platform environment: https://mitacs.crm.dynamics.com/", "Maker solution: https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions", "Model-driven app: https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f", "Power Pages site: https://enterprise-service-intake-hellox.powerappsportals.com", "Hidden ERP console: https://hellox.ca/esi/", "The portal is private for the interview tenant.", "Normal rule maintenance should use Routing Matrix.", "The clean final routing table should contain 81 active rules: 80 exact matrix rows plus one active fallback.", "No final solution export should be produced until the administrator requests it."], styles)
    story.append(Spacer(1, 2))
    story.append(Paragraph(f"Source Markdown: {SOURCE_PATH.relative_to(ROOT)}", styles["Note"]))

    document.build(story, onFirstPage=manual_pdf_page, onLaterPages=manual_pdf_page)


def build_document() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(
        doc,
        header_text="Enterprise Service Intake | User Manual",
        footer_text="Prepared for Mitacs hiring review - May 2026",
    )
    add_title_page(doc)
    add_quick_access(doc)
    add_portal(doc)
    add_internal_app(doc)
    add_routing_matrix(doc)
    add_monitoring(doc)
    add_accounts_notes(doc)
    doc.core_properties.title = "Enterprise Service Intake User Manual"
    doc.core_properties.subject = "Power Platform Senior Developer Take-Home Case"
    doc.core_properties.author = "Forrest Zhang"
    doc.core_properties.comments = "Generated from source-controlled project documentation."
    doc.save(DOCX_PATH)
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    build_document()
