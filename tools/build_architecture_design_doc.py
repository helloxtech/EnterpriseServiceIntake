#!/usr/bin/env python3
"""Build the Enterprise Service Intake architecture design DOCX artifact."""

from __future__ import annotations

import math
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "docs" / "submission"
ASSET_DIR = ROOT / "artifacts" / "doc-assets"
DOC_VERSION = "V3"
DOCX_PATH = SUBMISSION_DIR / "Enterprise_ServiceIntake_Architecture_Design_ForrestZhang_v3.docx"
PDF_PATH = SUBMISSION_DIR / "Enterprise_ServiceIntake_Architecture_Design_ForrestZhang_v3.pdf"

BLUE = "0067B1"
DARK_BLUE = "003B66"
LIGHT_BLUE = "EAF4FC"
GREY = "F2F4F7"
DARK_GREY = "344054"
MID_GREY = "667085"
GREEN = "167C3A"
AMBER = "B54708"
RED = "B42318"
INK = "101828"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D0D5DD", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


CellValue = str | list[str]


def set_cell_text(cell, text: CellValue, bold: bool = False, color: str = INK, size: int = 9) -> None:
    cell.text = ""
    if isinstance(text, list):
        for idx, item in enumerate(text):
            paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.left_indent = Pt(10)
            paragraph.paragraph_format.first_line_indent = Pt(-8)
            marker = paragraph.add_run("• ")
            marker.font.color.rgb = rgb(color)
            marker.font.size = Pt(size)
            run = paragraph.add_run(item)
            run.bold = bold
            run.font.color.rgb = rgb(color)
            run.font.size = Pt(size)
    else:
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.color.rgb = rgb(color)
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, header: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if header and row_idx == 0:
                set_cell_shading(cell, GREY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(DARK_BLUE)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    new_run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field_run(paragraph, label: str, value: str) -> None:
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = rgb(DARK_BLUE)
    paragraph.add_run(value)


def add_section_title(doc: Document, title: str, intro: str | None = None) -> None:
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if intro:
        paragraph = doc.add_paragraph(intro)
        paragraph.style = doc.styles["Body Text"]


def add_subtitle(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (2.2, 4.4)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=DARK_BLUE, size=9)
        set_cell_shading(cells[0], GREY)
        set_cell_text(cells[1], value, size=9)
    set_table_widths(table, list(widths))
    style_table(table, header=False)


def add_standard_table(
    doc: Document,
    headers: list[str],
    rows: list[list[CellValue]],
    widths: list[float],
    font_size: int = 8,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, size=font_size)
        set_cell_shading(table.rows[0].cells[idx], GREY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    set_table_widths(table, widths)
    style_table(table, header=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, typeface, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=typeface)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, typeface, fill: str) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, typeface, x2 - x1 - 34)
    line_height = typeface.size + 8 if hasattr(typeface, "size") else 24
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1 - total_height) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=typeface)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=typeface, fill=f"#{fill}")
        y += line_height


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str | None = None,
    fill: str = WHITE,
    outline: str = BLUE,
    title_color: str = DARK_BLUE,
) -> None:
    draw.rounded_rectangle(box, radius=20, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = box
    title_font = font(26, bold=True)
    subtitle_font = font(19)
    title_lines = wrap_text(draw, title, title_font, x2 - x1 - 40)
    y = y1 + 24
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, font=title_font, fill=f"#{title_color}")
        y += 34
    if subtitle:
        y += 4
        for line in wrap_text(draw, subtitle, subtitle_font, x2 - x1 - 42):
            bbox = draw.textbbox((0, 0), line, font=subtitle_font)
            draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, font=subtitle_font, fill=f"#{DARK_GREY}")
            y += 27


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = BLUE) -> None:
    draw.line([start, end], fill=f"#{color}", width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = math.pi / 7
    left = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    right = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, left, right], fill=f"#{color}")


def poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], color: str = BLUE, width: int = 4) -> None:
    if len(points) < 2:
        return
    for start, end in zip(points, points[1:]):
        draw.line([start, end], fill=f"#{color}", width=width)
    start = points[-2]
    end = points[-1]
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = math.pi / 7
    left = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    right = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, left, right], fill=f"#{color}")


def create_architecture_diagram(path: Path) -> None:
    image = PILImage.new("RGB", (2000, 1180), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 2000, 1180), fill="#FFFFFF")
    draw.text((70, 48), "Enterprise Service Intake - Runtime Architecture", font=font(42, bold=True), fill=f"#{DARK_BLUE}")
    draw.text((70, 104), "Customer intake, Dataverse rules, SharePoint documents, confirmation email, approval, and mock ERP sync", font=font(24), fill=f"#{MID_GREY}")
    draw.text((70, 138), "Arrowheads show runtime flow direction; colors indicate the owning platform area.", font=font(18), fill=f"#{MID_GREY}")

    boxes = {
        "customer": (70, 250, 350, 405),
        "portal": (450, 235, 750, 420),
        "dataverse": (850, 235, 1160, 420),
        "flow": (1260, 235, 1580, 420),
        "erp": (1670, 235, 1940, 420),
        "sharepoint": (450, 575, 750, 745),
        "plugin": (850, 575, 1160, 745),
        "rules": (1260, 575, 1580, 745),
        "app": (1260, 865, 1580, 1035),
    }
    draw_round_box(draw, boxes["customer"], "External Customer", "Authenticated portal user", fill=LIGHT_BLUE)
    draw_round_box(draw, boxes["portal"], "Power Pages", "Multi-step intake and live routing preview", fill=LIGHT_BLUE)
    draw_round_box(draw, boxes["dataverse"], "Dataverse", "Service Request, rules, logs, email activity", fill=WHITE, outline=DARK_BLUE)
    draw_round_box(draw, boxes["sharepoint"], "SharePoint", "Post-submit supporting files through document management", fill=WHITE, outline=GREEN, title_color=GREEN)
    draw_round_box(draw, boxes["plugin"], "C# Plugins", "Routing, SLA calculation, close guardrail", fill=GREY, outline=DARK_BLUE)
    draw_round_box(draw, boxes["rules"], "Editable Rule Matrix", "Routing rules, departments, SLA policies", fill=GREY, outline=DARK_BLUE)
    draw_round_box(draw, boxes["app"], "Model-driven App", "Coordinator queue and PCF status indicator", fill=WHITE, outline=GREEN, title_color=GREEN)
    draw_round_box(draw, boxes["flow"], "Power Automate", "Confirmation email, approval, HTTP POST, Try/Catch logging", fill=WHITE, outline=AMBER, title_color=AMBER)
    draw_round_box(draw, boxes["erp"], "Mock ERP API", "External ID written back to Dataverse", fill=WHITE, outline=RED, title_color=RED)

    arrow(draw, (350, 328), (450, 328), BLUE)
    arrow(draw, (750, 328), (850, 328), BLUE)
    arrow(draw, (1160, 328), (1260, 328), AMBER)
    arrow(draw, (1580, 328), (1670, 328), RED)
    arrow(draw, (600, 420), (600, 575), GREEN)
    arrow(draw, (1005, 420), (1005, 575), DARK_BLUE)
    arrow(draw, (1160, 660), (1260, 660), DARK_BLUE)
    poly_arrow(draw, [(1005, 745), (1005, 950), (1260, 950)], GREEN)

    legend_y = 1090
    legend = [
        (BLUE, "External intake"),
        (DARK_BLUE, "Transactional Dataverse logic"),
        (GREEN, "Document storage and internal operations"),
        (AMBER, "Email, approval, and integration orchestration"),
        (RED, "External system response"),
    ]
    x = 80
    for color, label in legend[:3]:
        draw.rounded_rectangle((x, legend_y, x + 28, legend_y + 28), radius=6, fill=f"#{color}")
        draw.text((x + 42, legend_y - 1), label, font=font(19), fill=f"#{DARK_GREY}")
        x += 545

    x = 80
    for color, label in legend[3:]:
        draw.rounded_rectangle((x, legend_y + 45, x + 28, legend_y + 73), radius=6, fill=f"#{color}")
        draw.text((x + 42, legend_y + 44), label, font=font(19), fill=f"#{DARK_GREY}")
        x += 545

    image.save(path)


def create_erd_diagram(path: Path) -> None:
    image = PILImage.new("RGB", (1900, 1350), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 48), "Dataverse Entity Relationship Diagram", font=font(42, bold=True), fill=f"#{DARK_BLUE}")
    draw.text((70, 104), "Core tables support request intake, routing, SLA assignment, SharePoint documents, integration logs, and error handling.", font=font(24), fill=f"#{MID_GREY}")
    draw.text((70, 138), "Arrowheads point to the dependent/many-side table; labels show relationship cardinality.", font=font(18), fill=f"#{MID_GREY}")

    def entity(box, title, fields, fill=WHITE, outline=BLUE):
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=18, fill=f"#{fill}", outline=f"#{outline}", width=3)
        draw.rectangle((x1, y1, x2, y1 + 54), fill=f"#{outline}")
        draw.text((x1 + 18, y1 + 13), title, font=font(24, bold=True), fill="white")
        y = y1 + 74
        for field in fields:
            draw.text((x1 + 18, y), "•", font=font(18, bold=True), fill=f"#{outline}")
            draw.text((x1 + 38, y), field, font=font(18), fill=f"#{INK}")
            y += 28

    entities = {
        "contact": (70, 235, 420, 385),
        "account": (70, 455, 420, 605),
        "category": (70, 735, 420, 910),
        "request": (650, 225, 1245, 635),
        "department": (1460, 235, 1830, 405),
        "sla": (1460, 455, 1830, 625),
        "rule": (1460, 735, 1830, 995),
        "document": (650, 740, 1245, 925),
        "sync": (650, 1075, 975, 1265),
        "error": (1075, 1075, 1400, 1265),
    }

    entity(entities["contact"], "Contact", ["External portal user", "Owns submitted requests"], LIGHT_BLUE)
    entity(entities["account"], "Account", ["Optional customer org", "Parent customer record"], LIGHT_BLUE)
    entity(entities["category"], "Service Category", ["Funding Agreement", "Technical Support", "Event Support"], LIGHT_BLUE)
    entity(
        entities["request"],
        "Service Request",
        [
            "Confirmation Number",
            "Title / Description",
            "Severity / Priority",
            "Lifecycle / Approval Status",
            "Assigned Department",
            "Applied SLA / Due Date",
            "External ERP ID",
            "Internal Resolution Notes",
        ],
        WHITE,
        DARK_BLUE,
    )
    entity(entities["department"], "Department", ["Finance", "IT Support", "Research Services"], GREY, DARK_BLUE)
    entity(entities["sla"], "SLA Policy", ["Response hours", "Business hours flag", "Escalation policy"], GREY, DARK_BLUE)
    entity(
        entities["rule"],
        "Routing Rule",
        [
            "Category match",
            "Severity / priority match",
            "Sort order",
            "Requires approval",
            "Requires documentation",
        ],
        GREY,
        DARK_BLUE,
    )
    entity(entities["document"], "SharePoint Documents", ["Stored by document management", "Linked to Service Request", "Optional review metadata"], WHITE, GREEN)
    entity(entities["sync"], "External Sync Log", ["Endpoint", "HTTP status", "Payload snapshot"], WHITE, AMBER)
    entity(entities["error"], "System Error Log", ["Source component", "Stage", "Correlation ID", "Technical detail"], WHITE, RED)

    def label(text: str, x: int, y: int, color: str = DARK_BLUE) -> None:
        bbox = draw.textbbox((0, 0), text, font=font(16, bold=True))
        draw.rounded_rectangle((x - 12, y - 14, x + (bbox[2] - bbox[0]) + 12, y + 18), radius=8, fill="white", outline=f"#{color}", width=1)
        draw.text((x, y - 11), text, font=font(16, bold=True), fill=f"#{color}")

    def rel(points: list[tuple[int, int]], color: str = DARK_BLUE) -> None:
        poly_arrow(draw, points, color=color, width=3)

    # Customer ownership and intake references.
    rel([(420, 310), (650, 310)], BLUE)
    label("1:N", 510, 286, BLUE)
    rel([(420, 530), (650, 530)], BLUE)
    label("1:N", 510, 506, BLUE)
    rel([(420, 822), (535, 822), (535, 585), (650, 585)], BLUE)
    label("1:N", 500, 795, BLUE)

    # Request lookups to operational routing targets.
    rel([(1460, 310), (1245, 310)], DARK_BLUE)
    label("1:N", 1328, 286, DARK_BLUE)
    rel([(1460, 530), (1245, 530)], DARK_BLUE)
    label("1:N", 1328, 506, DARK_BLUE)

    # Routing rule configuration relationships are routed outside the entity boxes.
    rel([(245, 910), (245, 1290), (1645, 1290), (1645, 995)], BLUE)
    label("1:N", 900, 1265, BLUE)
    rel([(1830, 310), (1870, 310), (1870, 805), (1830, 805)], DARK_BLUE)
    label("1:N", 1845, 555, DARK_BLUE)
    rel([(1830, 530), (1890, 530), (1890, 925), (1830, 925)], DARK_BLUE)
    label("1:N", 1865, 725, DARK_BLUE)

    # Child records and telemetry are routed around the document box to avoid crossings.
    rel([(948, 635), (948, 740)], GREEN)
    label("1:N", 970, 674, GREEN)
    rel([(760, 635), (760, 665), (575, 665), (575, 1170), (650, 1170)], AMBER)
    label("1:N", 585, 1028, AMBER)
    rel([(1135, 635), (1135, 665), (1425, 665), (1425, 1170), (1400, 1170)], RED)
    label("1:N", 1360, 955, RED)

    image.save(path)


def configure_styles(
    doc: Document,
    header_text: str = "Enterprise Service Intake | Architecture & Design Brief",
    footer_text: str = "Prepared for Mitacs hiring review - May 2026",
) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    body = styles["Body Text"]
    body.font.name = "Calibri"
    body.font.size = Pt(10.5)
    body.font.color.rgb = rgb(INK)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.space_after = Pt(3)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(header_text)
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MID_GREY)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(footer_text)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = rgb(MID_GREY)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Enterprise Service Intake")
    run.font.size = Pt(25)
    run.bold = True
    run.font.color.rgb = rgb(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run(f"Architecture & Design Brief - {DOC_VERSION}")
    sub_run.font.size = Pt(18)
    sub_run.bold = True
    sub_run.font.color.rgb = rgb(BLUE)

    meta_rows = [
        ("Candidate", "Forrest Zhang"),
        ("Role", "Power Platform Senior Developer - Take-Home Case"),
        ("Prepared For", "Mitacs hiring and technical review team"),
        ("Environment", "https://mitacs.crm.dynamics.com/"),
        ("Portal", "https://enterprise-service-intake-hellox.powerappsportals.com"),
        ("Status", f"Candidate submission brief - {DOC_VERSION}"),
    ]
    add_key_value_table(doc, meta_rows, widths=(1.7, 5.0))

    add_subtitle(doc, "Executive Summary")
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Body Text"]
    paragraph.add_run(
        "This candidate submission implements an enterprise-grade external service request intake process using Power Pages, "
        "Dataverse, Power Automate, C# plugins, and PCF. External customers submit authenticated multi-step "
        "requests, save drafts, receive real-time routing and SLA feedback, and upload required supporting documentation before final submission when the matched rule requires it. Dataverse "
        "holds the system of record while server-side plugins apply non-bypassable routing and close-validation "
        "rules. Power Automate handles applicant confirmation email, manager approval, mock ERP synchronization, and resilient error logging."
    )

    add_subtitle(doc, "End-to-End Outcome")
    add_numbered(
        doc,
        [
            "Customer submits a portal request and receives a formatted confirmation number.",
            "Customer can upload required or optional supporting files to the SharePoint document library associated to the saved Service Request.",
            "Dataverse plugin applies the matching routing/SLA rule and flags approval or documentation requirements.",
            "Internal coordinator reviews the request in the model-driven app with a PCF status indicator.",
            "Critical or high-priority items route to manager approval.",
            "Power Automate sends the confirmation number to the applicant and records failures in System Error Logs.",
            "Approved requests are posted to a mock REST endpoint and the external system ID is written back to Dataverse.",
            "Failures in approval or integration are captured in a custom System Error Log table.",
        ],
    )


def add_reviewer_links(doc: Document) -> None:
    add_section_title(doc, "Live Review Access", "These links and accounts are included so the hiring team can validate the submitted solution end to end during the review.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Area", "Value", "What Reviewers Can Validate"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, size=8)
        set_cell_shading(table.rows[0].cells[idx], GREY)
    rows = [
        ["Environment", "https://mitacs.crm.dynamics.com/", "Open Dataverse tables, model-driven app, flow runs, and solution components."],
        ["Maker Solution", "https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions", "Review solution-aware components and exports."],
        ["Model-driven App", "https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f", "Review request queue, form fields, and PCF coordinator experience."],
        ["Power Pages Site", "https://enterprise-service-intake-hellox.powerappsportals.com", "Submit a customer request, verify dynamic SLA/routing preview, save/resume drafts, and upload files."],
        ["Cloud Flows", "ESI - Send Confirmation Email; ESI - Approval and ERP Sync", "Review applicant email, approval, HTTP sync, run history, and Try/Catch error handling."],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=7)
    set_table_widths(table, [1.35, 2.85, 2.55])
    style_table(table, header=True)

    add_subtitle(doc, "Live Review Accounts")
    add_standard_table(
        doc,
        ["Account", "Purpose"],
        [
            ["forrest@hellosmart.ca", "Admin/customizer reviewer"],
            ["agent@hellosmart.ca", "Internal service coordinator"],
            ["manager@hellosmart.ca", "Approval manager"],
        ],
        [2.8, 3.9],
        font_size=9,
    )


def add_architecture(doc: Document, architecture_image: Path) -> None:
    add_section_title(doc, "Architecture Overview", "The solution keeps Dataverse as the authority for state while placing business logic in the component that best enforces the requirement.")
    doc.add_picture(str(architecture_image), width=Inches(6.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_subtitle(doc, "Component Placement")
    add_standard_table(
        doc,
        ["Requirement", "Component", "Reason"],
        [
            ["External intake", "Power Pages", ["Authenticated customer access", "Contact-scoped table permissions", "Multi-step forms", "Post-submit SharePoint upload", "Liquid/Web API extensibility"]],
            ["System of record", "Dataverse", ["Relational data model", "Security roles and ownership", "Auditing and solution packaging", "Consistent state across portal, app, flow, and API entry points"]],
            ["Routing and SLA", "C# plugin", ["Runs transactionally on create/update", "Applies the same rule result for all entry points", "Cannot be bypassed by imports, flows, APIs, or alternate clients"]],
            ["Close guardrail", "C# plugin", ["Enforces critical documentation requirements at the server boundary", "Stronger than form scripts or after-the-fact flow validation"]],
            ["Applicant email, approval, and ERP sync", "Power Automate", ["Sends the confirmation email", "Handles human approval", "Calls the mock ERP API", "Provides connector run history, retry visibility, and Try/Catch logging"]],
            ["Internal coordinator UX", "Model-driven app + PCF + Routing Matrix", ["Secure operational CRUD", "Coordinator queue views", "Editable rule matrix", "Compact visual severity, SLA, approval, and sync status"]],
        ],
        [1.45, 1.45, 3.85],
        font_size=7,
    )


def add_data_model(doc: Document, erd_image: Path) -> None:
    add_section_title(doc, "Dataverse Data Strategy", "The data model separates request state, rule configuration, documents, integration telemetry, and operational error records.")
    doc.add_picture(str(erd_image), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_subtitle(doc, "Core Tables")
    add_standard_table(
        doc,
        ["Table", "Purpose", "Key Relationships"],
        [
            ["Service Request", ["Primary intake and work-management row", "Stores confirmation number, status, priority, SLA, approval state, external ERP ID, and internal notes"], ["Contact", "Account", "Service Category", "Department", "SLA Policy", "SharePoint Documents", "External Sync Log", "System Error Log"]],
            ["Routing Rule", ["Configurable rules engine row", "Matches category, severity, and priority", "Supplies department, SLA, approval, and documentation requirements"], ["Service Category", "Department", "SLA Policy"]],
            ["Department", ["Internal routing destination", "Operational ownership for assigned requests"], ["Service Request", "Routing Rule"]],
            ["SLA Policy", ["Response target", "Business-hours and escalation policy assigned by routing"], ["Service Request", "Routing Rule"]],
            ["SharePoint Documents", ["Supporting files stored through Power Pages document management", "Files are linked to the saved Service Request after creation"], ["Service Request", "SharePoint Document Location"]],
            ["External Sync Log", ["Integration audit trail", "Tracks outbound ERP endpoint, status, and payload snapshot"], ["Service Request"]],
            ["System Error Log", ["Structured error capture", "Used by plugin, flow, and integration failure paths"], ["Service Request optional"]],
        ],
        [1.45, 3.25, 2.05],
        font_size=7,
    )

    add_subtitle(doc, "Confirmation Number")
    add_bullets(
        doc,
        [
            "Service Request uses a Dataverse autonumber format: SR-{yyyyMMdd}-{SEQNUM}.",
            "The number is generated server-side, remains stable after creation, and is safe to show externally.",
            "The formatted value gives customers a readable reference while preserving Dataverse row IDs for integration and support work.",
        ],
    )


def add_security(doc: Document) -> None:
    add_section_title(doc, "Security Strategy", "The design assumes least privilege across external portal users, internal coordinators, managers, and administrators.")
    add_subtitle(doc, "Power Pages Access")
    add_bullets(
        doc,
        [
            "External users authenticate to Power Pages and are associated to Contact rows.",
            "Table permissions are contact-scoped so users can create requests and view only requests tied to their contact.",
            "Reference data needed for dynamic preview, such as active service categories and routing outcomes, is exposed read-only.",
            "Portal forms exclude internal-only fields, integration payloads, internal notes, sync logs, and system error logs.",
        ],
    )

    add_subtitle(doc, "Internal Dataverse Access")
    add_bullets(
        doc,
        [
            "Coordinators use the model-driven app for operational request triage.",
            "Managers review approvals and can inspect high-priority items.",
            "Sensitive internal resolution notes are kept off the portal and should be protected with role and column-level security in production.",
            "Error and sync telemetry tables are intended for managers, admins, and support owners rather than broad staff access.",
        ],
    )

    add_subtitle(doc, "Security Tradeoffs")
    add_standard_table(
        doc,
        ["Area", "Lean Demo Choice", "Production Hardening"],
        [
            ["Portal identity", ["Authenticated Power Pages users in the interview tenant"], ["Use Entra External ID or configured B2C provider", "Apply lifecycle controls", "Add conditional access where appropriate"]],
            ["Uploads", ["Post-submit document-management page", "Files stored in SharePoint for the saved Service Request"], ["Apply file type limits", "Use malware scanning", "Review retention and DLP rules"]],
            ["Secrets", ["No tenant passwords or tokens committed to Git"], ["Use connection references and environment variables", "Use Key Vault-backed custom connectors", "Prefer managed identities where available"]],
        ],
        [1.35, 2.65, 2.75],
        font_size=7,
    )


def add_portal_ux(doc: Document) -> None:
    add_section_title(doc, "External User Experience", "The portal is scoped as a clean customer intake experience rather than an internal operations surface.")
    add_subtitle(doc, "Submission Flow")
    add_numbered(
        doc,
        [
            "User signs in to the Power Pages site.",
            "User starts a new service request, enters request details, and can save a draft.",
            "User selects category, impact, and urgency.",
            "Portal dynamically previews expected department, SLA target, approval requirement, and documentation requirement before final submission.",
            "When documentation is required, the portal saves the Draft request and holds the user in the Files step until a file is uploaded.",
            "User reviews and submits the request, then receives a confirmation number.",
            "Optional files can be added after submission through the same secure request-specific upload path.",
        ],
    )
    add_subtitle(doc, "Dynamic Preview Implementation")
    add_bullets(
        doc,
        [
            "The portal reads eligible routing/SLA data through Power Pages Web API and/or Liquid-rendered data.",
            "Client-side JavaScript recalculates the displayed routing destination and expected SLA when the user changes category, impact, or urgency.",
            "The preview avoids a full page reload and mirrors the server-side plugin rule result, while the plugin remains the authoritative enforcement layer.",
        ],
    )


def add_automation(doc: Document) -> None:
    add_section_title(doc, "Automation And Integration", "The cloud flow handles long-running human approval and integration work that should remain visible in run history.")
    add_subtitle(doc, "Flow: ESI - Approval and ERP Sync")
    add_standard_table(
        doc,
        ["Stage", "Implementation"],
        [
            ["Trigger", ["Dataverse row added or modified", "Approval required", "Approval pending", "Integration unsynced"]],
            ["Try scope", ["Start and wait for manager approval", "Branch on approval decision", "Call mock ERP API", "Update Service Request", "Write External Sync Log"]],
            ["HTTP sync", ["POST approved request details to https://hellox.ca/api/mock/enterprise-service-intake/erp", "Store returned external ID in Dataverse"]],
            ["Reject branch", ["Mark approval rejected", "Skip ERP sync"]],
            ["Catch scope", ["Runs after failure, timeout, or skipped Try scope", "Writes System Error Log with flow run correlation details", "Marks request failed where appropriate"]],
        ],
        [1.35, 5.4],
        font_size=8,
    )
    add_subtitle(doc, "Flow: ESI - Send Confirmation Email")
    add_standard_table(
        doc,
        ["Stage", "Implementation"],
        [
            ["Trigger", ["Dataverse row added or modified", "Lifecycle status is Submitted"]],
            ["Try scope", ["Validate applicant Contact and email", "Send HTML confirmation through Office 365 Outlook Send an email (V2)", "Include confirmation number, request title, and submitted timestamp", "Update customer-visible notes"]],
            ["Skip path", ["Missing Contact or email writes a System Error Log", "Flow avoids silent failures"]],
            ["Catch scope", ["Runs after failure, timeout, or skipped Try scope", "Writes System Error Log with flow run correlation details"]],
        ],
        [1.35, 5.4],
        font_size=8,
    )
    add_subtitle(doc, "Resiliency Pattern")
    add_bullets(
        doc,
        [
            "Try/Catch scopes make failures observable and reviewable during the live demo.",
            "The custom System Error Log captures source component, stage, message, technical detail, correlation ID, and related request.",
            "Request integration status is updated to failed when the integration path cannot complete.",
            "Office 365 Outlook action output, Flow run history, and Approval records provide proof if tenant email delivery is restricted.",
        ],
    )


def add_pro_code(doc: Document) -> None:
    add_section_title(doc, "Pro-Code Extensibility", "The pro-code pieces are intentionally placed where low-code would either be bypassable or less effective.")
    add_standard_table(
        doc,
        ["Component", "Location", "Responsibility", "Why Here"],
        [
            ["ServiceRequestRoutingPlugin", "src/plugins/ServiceIntake.Plugins/ServiceRequestRoutingPlugin.cs", ["Evaluates routing rules", "Assigns department and SLA", "Sets due date and lifecycle defaults", "Sets approval and documentation requirements"], ["PreOperation server-side execution", "Transactional rule result", "Consistent across portal, model app, imports, flows, and APIs"]],
            ["ServiceRequestClosureGuardPlugin", "src/plugins/ServiceIntake.Plugins/ServiceRequestClosureGuardPlugin.cs", ["Blocks critical request closure", "Requires internal resolution notes", "Requires documentation evidence"], ["Server-side validation", "Prevents bypass from forms, imports, flows, and APIs"]],
            ["SlaStatusIndicator PCF", "src/pcf/SlaStatusIndicator/", ["Displays severity", "Shows SLA, approval, and sync status", "Keeps coordinator status scanning compact"], ["Improves internal model-driven usability", "Does not duplicate authoritative business logic"]],
            ["Provisioning Utility", "src/scripts/ServiceIntake.Provisioning/", ["Creates metadata", "Registers plugins", "Provisions app, forms, views, dashboards, and sample data", "Patches flow definitions"], ["Repeatable environment build", "Reviewable ALM steps"]],
        ],
        [1.45, 1.75, 2.05, 1.5],
        font_size=6,
    )


def add_alm_verification(doc: Document) -> None:
    add_section_title(doc, "ALM, Verification, And Demo Evidence", "The repository contains both deployable solution packages and unpacked source so reviewers can inspect the implementation.")
    add_subtitle(doc, "Deliverables")
    add_standard_table(
        doc,
        ["Artifact", "Location"],
        [
            ["Managed solution", "solution/export/Enterprise_ServiceIntake_ForrestZhang_managed.zip"],
            ["Unmanaged solution", "solution/export/Enterprise_ServiceIntake_ForrestZhang_unmanaged.zip"],
            ["Unpacked solution source", "solution/unpacked/managed/ and solution/unpacked/unmanaged/"],
            ["C# plugin source", "src/plugins/ServiceIntake.Plugins/"],
            ["PCF source", "src/pcf/SlaStatusIndicator/"],
            ["Power Pages source", "src/powerpages/ and powerpages-live/"],
            ["Demo script", "docs/demo/demo-script.md"],
            ["Verification notes", "docs/evidence/verification.md"],
        ],
        [2.15, 4.6],
        font_size=8,
    )

    add_subtitle(doc, "Verified Behavior")
    add_standard_table(
        doc,
        ["Check", "Evidence"],
        [
            ["Plugin build", "Passed."],
            ["Provisioning utility", ["Builds with zero warnings", "NuGet vulnerability audit reports no vulnerable packages"]],
            ["PCF", "Build and pac pcf push passed."],
            ["Power Pages", ["Dynamic preview shows Finance", "4 hour SLA and approval required display before submit", "Post-submit SharePoint upload is available from the confirmation dialog"]],
            ["Portal submission", ["Portal demo request created with formatted confirmation SR-20260521-001004", "Outlook confirmation smoke test created SR-20260521-001018"]],
            ["Closure guard", ["Blocked undocumented critical closure", "Allowed documented closure"]],
            ["Flows", ["Approval/ERP flow is active and solution-aware", "Confirmation email flow is active and solution-aware", "Both use Try/Catch logging patterns"]],
            ["Solution export/unpack", "Managed and unmanaged exports and unpacked source generated successfully."],
        ],
        [1.55, 5.2],
        font_size=7,
    )

    add_subtitle(doc, "Known Demo Notes")
    add_bullets(
        doc,
        [
            "If email delivery is restricted, use Office 365 Outlook action output, Approval records, and Flow run history as evidence.",
            "The HelloX mock ERP endpoint avoids third-party API keys and includes a deliberate failure mode for Catch-scope demos.",
            "The hidden HelloX console at https://hellox.ca/esi/ is noindex and not linked from public navigation.",
            "If a target environment does not automatically bind the PCF control on import, add the PCF control to the coordinator form field in the form designer.",
        ],
    )


def add_appendix(doc: Document) -> None:
    add_section_title(doc, "Live Demonstration Path", "The demo is designed to show the full lifecycle and then drill into implementation decisions.")
    add_numbered(
        doc,
        [
            "Walk through the ERD and architecture choices.",
            "Submit a critical funding request from the portal and show dynamic preview before submit.",
            "Open the request in the model-driven app and review routing, SLA, approval, sync status, dashboards, and internal fields.",
            "Open the cloud flow and explain approval, HTTP POST, writeback, sync log, reject branch, and Catch scope.",
            "Demonstrate or explain failure handling through System Error Log rows.",
            "Demonstrate the critical close guardrail through the smoke test or form behavior.",
            "Show managed solution, unpacked source, PCF, plugin source, Power Pages source, and GitHub repository.",
        ],
    )

    add_subtitle(doc, "Prepared Q&A")
    add_standard_table(
        doc,
        ["Question", "Answer"],
        [
            ["Why plugin over flow for routing?", "Routing must be transactional and consistent for every entry point, including portal, model app, imports, flows, and APIs."],
            ["Why plugin for close validation?", "The requirement says agents cannot bypass documentation requirements; server-side PreOperation validation is the correct enforcement point."],
            ["Why flow for approvals and ERP sync?", "Approvals and integration orchestration need human workflow, connector history, retry visibility, and run evidence."],
            ["How do customers only see their data?", "Power Pages users are tied to Contacts, and contact-scoped table permissions restrict request visibility to owned rows."],
            ["How is this deployable?", "Components are solution-aware, exported as managed/unmanaged zips, unpacked with PAC, and source-controlled with raw plugin and PCF code."],
        ],
        [2.3, 4.45],
        font_size=7,
    )


def pdf_color(hex_value: str) -> colors.HexColor:
    return colors.HexColor(f"#{hex_value}")


def create_pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "ESITitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=pdf_color(DARK_BLUE),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "Subtitle": ParagraphStyle(
            "ESISubtitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=pdf_color(BLUE),
            spaceAfter=14,
        ),
        "H1": ParagraphStyle(
            "ESIH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=pdf_color(BLUE),
            spaceBefore=12,
            spaceAfter=7,
        ),
        "H2": ParagraphStyle(
            "ESIH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=15,
            textColor=pdf_color(DARK_BLUE),
            spaceBefore=9,
            spaceAfter=5,
        ),
        "Body": ParagraphStyle(
            "ESIBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=12.2,
            textColor=pdf_color(INK),
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "Small": ParagraphStyle(
            "ESISmall",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.2,
            leading=9.2,
            textColor=pdf_color(INK),
            wordWrap="CJK",
        ),
        "Header": ParagraphStyle(
            "ESIHeader",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=7.5,
            leading=9,
            textColor=pdf_color(DARK_BLUE),
            wordWrap="CJK",
        ),
    }


def pp(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def pdf_cell(value: CellValue, style: ParagraphStyle) -> Paragraph:
    if isinstance(value, list):
        bullet_html = "<br/>".join(f"&#8226; {escape(item)}" for item in value)
        return Paragraph(bullet_html, style)
    return Paragraph(escape(str(value)), style)


def pdf_heading(story: list, text: str, styles: dict[str, ParagraphStyle], level: int = 1) -> None:
    story.append(pp(text, styles["H1" if level == 1 else "H2"]))


def pdf_bullets(story: list, items: list[str], styles: dict[str, ParagraphStyle]) -> None:
    story.append(
        ListFlowable(
            [ListItem(pp(item, styles["Body"]), leftIndent=12) for item in items],
            bulletType="bullet",
            leftIndent=16,
            bulletFontName="Helvetica",
            bulletFontSize=8,
        )
    )
    story.append(Spacer(1, 4))


def pdf_numbers(story: list, items: list[str], styles: dict[str, ParagraphStyle]) -> None:
    story.append(
        ListFlowable(
            [ListItem(pp(item, styles["Body"]), leftIndent=14) for item in items],
            bulletType="1",
            leftIndent=18,
            bulletFontName="Helvetica",
            bulletFontSize=8,
        )
    )
    story.append(Spacer(1, 4))


def pdf_table(
    story: list,
    headers: list[str],
    rows: list[list[CellValue]],
    widths: list[float],
    styles: dict[str, ParagraphStyle],
    font_style: str = "Small",
) -> None:
    header_style = styles["Header"]
    cell_style = styles[font_style]
    data = [[pp(header, header_style) for header in headers]]
    data.extend([[pdf_cell(value, cell_style) for value in row] for row in rows])
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), pdf_color(GREY)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def pdf_key_value_table(story: list, rows: list[tuple[str, str]], styles: dict[str, ParagraphStyle]) -> None:
    data = []
    for label, value in rows:
        data.append([pp(label, styles["Header"]), pp(value, styles["Small"])])
    table = Table(data, colWidths=[1.6 * inch, 5.1 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), pdf_color(GREY)),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def pdf_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, letter[0], letter[1], stroke=0, fill=1)
    canvas.setStrokeColor(pdf_color(BLUE))
    canvas.setLineWidth(0.6)
    canvas.line(0.75 * inch, 10.33 * inch, 7.75 * inch, 10.33 * inch)
    canvas.setFillColor(pdf_color(MID_GREY))
    canvas.setFont("Helvetica", 7.5)
    canvas.drawRightString(7.75 * inch, 10.43 * inch, "Enterprise Service Intake | Architecture & Design Brief")
    canvas.drawString(0.75 * inch, 0.45 * inch, "Prepared for Mitacs hiring review - May 2026")
    canvas.drawRightString(7.75 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(architecture_image: Path, erd_image: Path) -> None:
    styles = create_pdf_styles()
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.82 * inch,
        bottomMargin=0.72 * inch,
        title="Enterprise Service Intake Architecture & Design Brief",
        author="Forrest Zhang",
        subject="Power Platform Senior Developer Take-Home Case",
    )
    story: list = []

    story.append(pp("Enterprise Service Intake", styles["Title"]))
    story.append(pp(f"Architecture & Design Brief - {DOC_VERSION}", styles["Subtitle"]))
    pdf_key_value_table(
        story,
        [
            ("Candidate", "Forrest Zhang"),
            ("Role", "Power Platform Senior Developer - Take-Home Case"),
            ("Prepared For", "Mitacs hiring and technical review team"),
            ("Environment", "https://mitacs.crm.dynamics.com/"),
            ("Portal", "https://enterprise-service-intake-hellox.powerappsportals.com"),
            ("Status", f"Candidate submission brief - {DOC_VERSION}"),
        ],
        styles,
    )
    pdf_heading(story, "Executive Summary", styles, level=2)
    story.append(
        pp(
            "This candidate submission implements an enterprise-grade external service request intake process using Power Pages, Dataverse, Power Automate, C# plugins, and PCF. External customers submit authenticated multi-step requests, save drafts, receive real-time routing and SLA feedback, and upload required supporting documentation before final submission when the matched rule requires it. Dataverse holds the system of record while server-side plugins apply non-bypassable routing and close-validation rules. Power Automate handles applicant confirmation email, manager approval, mock ERP synchronization, and resilient error logging.",
            styles["Body"],
        )
    )
    pdf_heading(story, "End-to-End Outcome", styles, level=2)
    pdf_numbers(
        story,
        [
            "Customer submits a portal request and receives a formatted confirmation number.",
            "Customer can upload required or optional supporting files to the SharePoint document library associated to the saved Service Request.",
            "Dataverse plugin applies the matching routing/SLA rule and flags approval or documentation requirements.",
            "Internal coordinator reviews the request in the model-driven app with a PCF status indicator.",
            "Critical or high-priority items route to manager approval.",
            "Power Automate sends the confirmation number to the applicant and records failures in System Error Logs.",
            "Approved requests are posted to a mock REST endpoint and the external system ID is written back to Dataverse.",
            "Failures in approval or integration are captured in a custom System Error Log table.",
        ],
        styles,
    )
    story.append(PageBreak())

    pdf_heading(story, "Live Review Access", styles)
    story.append(pp("These links and accounts are included so the hiring team can validate the submitted solution end to end during the review.", styles["Body"]))
    pdf_table(
        story,
        ["Area", "Value", "What Reviewers Can Validate"],
        [
            ["Environment", "https://mitacs.crm.dynamics.com/", "Open Dataverse tables, model-driven app, flow runs, and solution components."],
            ["Maker Solution", "https://make.powerapps.com/environments/99dd50ed-a753-e37f-912c-78a022b12b09/solutions", "Review solution-aware components and exports."],
            ["Model-driven App", "https://mitacs.crm.dynamics.com/main.aspx?appid=3de4f813-b454-f111-bec7-000d3a3aca8f", "Review request queue, form fields, and PCF coordinator experience."],
            ["Power Pages Site", "https://enterprise-service-intake-hellox.powerappsportals.com", "Submit a customer request, verify dynamic SLA/routing preview, save/resume drafts, and upload files."],
            ["Cloud Flows", "ESI - Send Confirmation Email; ESI - Approval and ERP Sync", "Review applicant email, approval, HTTP sync, run history, and Try/Catch error handling."],
        ],
        [1.25, 2.75, 2.75],
        styles,
    )
    pdf_heading(story, "Live Review Accounts", styles, level=2)
    pdf_table(
        story,
        ["Account", "Purpose"],
        [
            ["forrest@hellosmart.ca", "Admin/customizer reviewer"],
            ["agent@hellosmart.ca", "Internal service coordinator"],
            ["manager@hellosmart.ca", "Approval manager"],
        ],
        [2.8, 3.9],
        styles,
    )
    pdf_heading(story, "Architecture Overview", styles)
    story.append(pp("The solution keeps Dataverse as the authority for state while placing business logic in the component that best enforces the requirement.", styles["Body"]))
    story.append(RLImage(str(architecture_image), width=6.7 * inch, height=4.17 * inch))
    story.append(Spacer(1, 8))
    pdf_table(
        story,
        ["Requirement", "Component", "Reason"],
        [
            ["External intake", "Power Pages", ["Authenticated customer access", "Contact-scoped permissions", "Multi-step forms", "Post-submit SharePoint upload", "Liquid/Web API extensibility"]],
            ["System of record", "Dataverse", ["Relational model", "Security roles and ownership", "Auditing and solution packaging", "Consistent state across portal, app, flow, and APIs"]],
            ["Routing and SLA", "C# plugin", ["Runs transactionally on create/update", "Same rule result for all entry points", "Cannot be bypassed by imports, flows, APIs, or alternate clients"]],
            ["Close guardrail", "C# plugin", ["Enforces critical documentation requirements", "Runs at the server boundary"]],
            ["Applicant email, approval, and ERP sync", "Power Automate", ["Sends the confirmation email", "Handles human approval", "Calls the mock ERP API", "Provides run history, retry visibility, and Try/Catch logging"]],
            ["Internal coordinator UX", "Model-driven app + PCF + Routing Matrix", ["Secure operational CRUD", "Coordinator queue views", "Editable rule matrix", "Compact visual severity, SLA, approval, and sync status"]],
        ],
        [1.4, 1.4, 3.95],
        styles,
    )
    story.append(PageBreak())

    pdf_heading(story, "Dataverse Data Strategy", styles)
    story.append(pp("The data model separates request state, rule configuration, documents, integration telemetry, and operational error records.", styles["Body"]))
    story.append(RLImage(str(erd_image), width=6.7 * inch, height=4.69 * inch))
    story.append(Spacer(1, 8))
    pdf_table(
        story,
        ["Table", "Purpose", "Key Relationships"],
        [
            ["Service Request", ["Primary intake and work-management row", "Stores confirmation number, status, priority, SLA, approval state, external ERP ID, and internal notes"], ["Contact", "Account", "Service Category", "Department", "SLA Policy", "SharePoint Documents", "External Sync Log", "System Error Log"]],
            ["Routing Rule", ["Configurable rules engine row", "Matches category, severity, and priority", "Supplies department, SLA, approval, and documentation requirements"], ["Service Category", "Department", "SLA Policy"]],
            ["Department", ["Internal routing destination", "Operational ownership for assigned requests"], ["Service Request", "Routing Rule"]],
            ["SLA Policy", ["Response target", "Business-hours and escalation policy assigned by routing"], ["Service Request", "Routing Rule"]],
            ["SharePoint Documents", ["Supporting files stored through Power Pages document management", "Files linked to the saved Service Request after creation"], ["Service Request", "SharePoint Document Location"]],
            ["External Sync Log", ["Integration audit trail", "Tracks outbound ERP endpoint, status, and payload snapshot"], ["Service Request"]],
            ["System Error Log", ["Structured error capture", "Used by plugin, flow, and integration failure paths"], ["Service Request optional"]],
        ],
        [1.35, 3.2, 2.2],
        styles,
    )
    pdf_heading(story, "Confirmation Number", styles, level=2)
    pdf_bullets(
        story,
        [
            "Service Request uses a Dataverse autonumber format: SR-{yyyyMMdd}-{SEQNUM}.",
            "The number is generated server-side, remains stable after creation, and is safe to show externally.",
            "The formatted value gives customers a readable reference while preserving Dataverse row IDs for integration and support work.",
        ],
        styles,
    )
    pdf_heading(story, "Security Strategy", styles)
    story.append(pp("The design assumes least privilege across external portal users, internal coordinators, managers, and administrators.", styles["Body"]))
    pdf_heading(story, "Power Pages Access", styles, level=2)
    pdf_bullets(
        story,
        [
            "External users authenticate to Power Pages and are associated to Contact rows.",
            "Table permissions are contact-scoped so users can create requests and view only requests tied to their contact.",
            "Reference data needed for dynamic preview is exposed read-only.",
            "Portal forms exclude internal-only fields, integration payloads, internal notes, sync logs, and system error logs.",
        ],
        styles,
    )
    pdf_heading(story, "Internal Dataverse Access", styles, level=2)
    pdf_bullets(
        story,
        [
            "Coordinators use the model-driven app for operational request triage.",
            "Managers review approvals and can inspect high-priority items.",
            "Sensitive internal resolution notes are kept off the portal and should be protected with role and column-level security in production.",
            "Error and sync telemetry tables are intended for managers, admins, and support owners rather than broad staff access.",
        ],
        styles,
    )
    pdf_heading(story, "Security Tradeoffs", styles, level=2)
    pdf_table(
        story,
        ["Area", "Lean Demo Choice", "Production Hardening"],
        [
            ["Portal identity", ["Authenticated Power Pages users in the interview tenant"], ["Use Entra External ID or configured B2C provider", "Apply lifecycle controls", "Add conditional access where appropriate"]],
            ["Uploads", ["Post-submit document-management page", "Files stored in SharePoint for the saved Service Request"], ["Apply file type limits", "Use malware scanning", "Review retention and DLP rules"]],
            ["Secrets", ["No tenant passwords or tokens committed to Git"], ["Use connection references and environment variables", "Use Key Vault-backed custom connectors", "Prefer managed identities where available"]],
        ],
        [1.25, 2.7, 2.8],
        styles,
    )

    pdf_heading(story, "External User Experience", styles)
    story.append(pp("The portal is scoped as a clean customer intake experience rather than an internal operations surface.", styles["Body"]))
    pdf_heading(story, "Submission Flow", styles, level=2)
    pdf_numbers(
        story,
        [
            "User signs in to the Power Pages site.",
            "User starts a new service request, enters request details, and can save a draft.",
            "User selects category, impact, and urgency.",
            "Portal dynamically previews expected department, SLA target, approval requirement, and documentation requirement before final submission.",
            "When documentation is required, the portal saves the Draft request and holds the user in the Files step until a file is uploaded.",
            "User reviews and submits the request, then receives a confirmation number.",
            "Optional files can be added after submission through the same secure request-specific upload path.",
        ],
        styles,
    )
    pdf_heading(story, "Dynamic Preview Implementation", styles, level=2)
    pdf_bullets(
        story,
        [
            "The portal reads eligible routing/SLA data through Power Pages Web API and/or Liquid-rendered data.",
            "Client-side JavaScript recalculates the displayed routing destination and expected SLA when the user changes category, impact, or urgency.",
            "The preview avoids a full page reload and mirrors the server-side plugin rule result, while the plugin remains the authoritative enforcement layer.",
        ],
        styles,
    )
    pdf_heading(story, "Automation And Integration", styles)
    story.append(pp("The cloud flow handles long-running human approval and integration work that should remain visible in run history.", styles["Body"]))
    pdf_table(
        story,
        ["Stage", "Implementation"],
        [
            ["Trigger", ["Dataverse row added or modified", "Approval required", "Approval pending", "Integration unsynced"]],
            ["Try scope", ["Start and wait for manager approval", "Branch on approval decision", "Call mock ERP API", "Update Service Request", "Write External Sync Log"]],
            ["HTTP sync", ["POST approved request details to https://hellox.ca/api/mock/enterprise-service-intake/erp", "Store returned external ID in Dataverse"]],
            ["Reject branch", ["Mark approval rejected", "Skip ERP sync"]],
            ["Catch scope", ["Runs after failure, timeout, or skipped Try scope", "Writes System Error Log with flow run correlation details", "Marks request failed where appropriate"]],
        ],
        [1.25, 5.5],
        styles,
    )
    pdf_heading(story, "Flow: ESI - Send Confirmation Email", styles, level=2)
    pdf_table(
        story,
        ["Stage", "Implementation"],
        [
            ["Trigger", ["Dataverse row added or modified", "Lifecycle status is Submitted"]],
            ["Try scope", ["Validate applicant Contact and email", "Send HTML confirmation through Office 365 Outlook Send an email (V2)", "Include confirmation number, request title, and submitted timestamp", "Update customer-visible notes"]],
            ["Skip path", ["Missing Contact or email writes a System Error Log", "Flow avoids silent failures"]],
            ["Catch scope", ["Runs after failure, timeout, or skipped Try scope", "Writes System Error Log with flow run correlation details"]],
        ],
        [1.25, 5.5],
        styles,
    )
    pdf_heading(story, "Resiliency Pattern", styles, level=2)
    pdf_bullets(
        story,
        [
            "Try/Catch scopes make failures observable and reviewable during the live demo.",
            "The custom System Error Log captures source component, stage, message, technical detail, correlation ID, and related request.",
            "Request integration status is updated to failed when the integration path cannot complete.",
            "Office 365 Outlook action output, Flow run history, and Approval records provide proof if tenant email delivery is restricted.",
        ],
        styles,
    )

    pdf_heading(story, "Pro-Code Extensibility", styles)
    story.append(pp("The pro-code pieces are intentionally placed where low-code would either be bypassable or less effective.", styles["Body"]))
    pdf_table(
        story,
        ["Component", "Location", "Responsibility", "Why Here"],
        [
            ["ServiceRequestRoutingPlugin", "src/plugins/ServiceIntake.Plugins/ServiceRequestRoutingPlugin.cs", ["Evaluates routing rules", "Assigns department and SLA", "Sets due date and lifecycle defaults", "Sets approval and documentation requirements"], ["PreOperation server-side execution", "Transactional rule result", "Consistent across portal, model app, imports, flows, and APIs"]],
            ["ServiceRequestClosureGuardPlugin", "src/plugins/ServiceIntake.Plugins/ServiceRequestClosureGuardPlugin.cs", ["Blocks critical request closure", "Requires internal resolution notes", "Requires documentation evidence"], ["Server-side validation", "Prevents bypass from forms, imports, flows, and APIs"]],
            ["SlaStatusIndicator PCF", "src/pcf/SlaStatusIndicator/", ["Displays severity", "Shows SLA, approval, and sync status", "Keeps coordinator status scanning compact"], ["Improves internal model-driven usability", "Does not duplicate authoritative business logic"]],
            ["Provisioning Utility", "src/scripts/ServiceIntake.Provisioning/", ["Creates metadata", "Registers plugins", "Provisions app, forms, views, dashboards, and sample data", "Patches flow definitions"], ["Repeatable environment build", "Reviewable ALM steps"]],
        ],
        [1.35, 1.55, 2.05, 1.8],
        styles,
    )
    story.append(PageBreak())

    pdf_heading(story, "ALM, Verification, And Demo Evidence", styles)
    story.append(pp("The repository contains both deployable solution packages and unpacked source so reviewers can inspect the implementation.", styles["Body"]))
    pdf_heading(story, "Deliverables", styles, level=2)
    pdf_table(
        story,
        ["Artifact", "Location"],
        [
            ["Managed solution", "solution/export/Enterprise_ServiceIntake_ForrestZhang_managed.zip"],
            ["Unmanaged solution", "solution/export/Enterprise_ServiceIntake_ForrestZhang_unmanaged.zip"],
            ["Unpacked solution source", "solution/unpacked/managed/ and solution/unpacked/unmanaged/"],
            ["C# plugin source", "src/plugins/ServiceIntake.Plugins/"],
            ["PCF source", "src/pcf/SlaStatusIndicator/"],
            ["Power Pages source", "src/powerpages/ and powerpages-live/"],
            ["Demo script", "docs/demo/demo-script.md"],
            ["Verification notes", "docs/evidence/verification.md"],
        ],
        [2.1, 4.65],
        styles,
    )
    pdf_heading(story, "Verified Behavior", styles, level=2)
    pdf_table(
        story,
        ["Check", "Evidence"],
        [
            ["Plugin build", "Passed."],
            ["Provisioning utility", ["Builds with zero warnings", "NuGet vulnerability audit reports no vulnerable packages"]],
            ["PCF", "Build and pac pcf push passed."],
            ["Power Pages", ["Dynamic preview shows Finance", "4 hour SLA and approval required display before submit", "Post-submit SharePoint upload is available from the confirmation dialog"]],
            ["Portal submission", ["Portal demo request created with formatted confirmation SR-20260521-001004", "Outlook confirmation smoke test created SR-20260521-001018"]],
            ["Closure guard", ["Blocked undocumented critical closure", "Allowed documented closure"]],
            ["Flows", ["Approval/ERP flow is active and solution-aware", "Confirmation email flow is active and solution-aware", "Both use Try/Catch logging patterns"]],
            ["Solution export/unpack", "Managed and unmanaged exports and unpacked source generated successfully."],
        ],
        [1.55, 5.2],
        styles,
    )
    pdf_heading(story, "Known Demo Notes", styles, level=2)
    pdf_bullets(
        story,
        [
            "If email delivery is restricted, use Office 365 Outlook action output, Approval records, and Flow run history as evidence.",
            "The HelloX mock ERP endpoint avoids third-party API keys and includes a deliberate failure mode for Catch-scope demos.",
            "The hidden HelloX console at https://hellox.ca/esi/ is noindex and not linked from public navigation.",
            "If a target environment does not automatically bind the PCF control on import, add the PCF control to the coordinator form field in the form designer.",
        ],
        styles,
    )

    pdf_heading(story, "Live Demonstration Path", styles)
    story.append(pp("The demo is designed to show the full lifecycle and then drill into implementation decisions.", styles["Body"]))
    pdf_numbers(
        story,
        [
            "Walk through the ERD and architecture choices.",
            "Submit a critical funding request from the portal and show dynamic preview before submit.",
            "Open the request in the model-driven app and review routing, SLA, approval, sync status, dashboards, and internal fields.",
            "Open the cloud flow and explain approval, HTTP POST, writeback, sync log, reject branch, and Catch scope.",
            "Demonstrate or explain failure handling through System Error Log rows.",
            "Demonstrate the critical close guardrail through the smoke test or form behavior.",
            "Show managed solution, unpacked source, PCF, plugin source, Power Pages source, and GitHub repository.",
        ],
        styles,
    )
    pdf_heading(story, "Prepared Q&A", styles, level=2)
    pdf_table(
        story,
        ["Question", "Answer"],
        [
            ["Why plugin over flow for routing?", "Routing must be transactional and consistent for every entry point, including portal, model app, imports, flows, and APIs."],
            ["Why plugin for close validation?", "The requirement says agents cannot bypass documentation requirements; server-side PreOperation validation is the correct enforcement point."],
            ["Why flow for approvals and ERP sync?", "Approvals and integration orchestration need human workflow, connector history, retry visibility, and run evidence."],
            ["How do customers only see their data?", "Power Pages users are tied to Contacts, and contact-scoped table permissions restrict request visibility to owned rows."],
            ["How is this deployable?", "Components are solution-aware, exported as managed/unmanaged zips, unpacked with PAC, and source-controlled with raw plugin and PCF code."],
        ],
        [2.25, 4.5],
        styles,
    )

    document.build(story, onFirstPage=pdf_page, onLaterPages=pdf_page)


def build_document() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_image = ASSET_DIR / "architecture-overview-v3.png"
    erd_image = ASSET_DIR / "dataverse-erd-v3.png"
    create_architecture_diagram(architecture_image)
    create_erd_diagram(erd_image)

    doc = Document()
    configure_styles(doc)
    add_title_page(doc)
    add_reviewer_links(doc)
    add_architecture(doc, architecture_image)
    add_data_model(doc, erd_image)
    add_security(doc)
    add_portal_ux(doc)
    add_automation(doc)
    add_pro_code(doc)
    add_alm_verification(doc)
    add_appendix(doc)
    doc.core_properties.title = "Enterprise Service Intake Architecture & Design Brief"
    doc.core_properties.subject = "Power Platform Senior Developer Take-Home Case"
    doc.core_properties.author = "Forrest Zhang"
    doc.core_properties.comments = "Generated from source-controlled project documentation."
    doc.save(DOCX_PATH)
    build_pdf(architecture_image, erd_image)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    build_document()
